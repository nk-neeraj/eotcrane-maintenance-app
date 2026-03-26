import streamlit as st
import pandas as pd
import plotly.express as px
from datetime import datetime
import db as db
import sqlite3
import io
from docx import Document
from docx.shared import RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Page Configuration
st.set_page_config(page_title="EOT Crane Maintenance Tracking System", layout="wide", page_icon="🏗️")

# Custom CSS for styling
st.markdown("""
<style>
    .overdue {
        color: white !important;
        background-color: #ff4b4b !important;
        padding: 5px;
        border-radius: 5px;
        font-weight: bold;
    }
    .warning {
        color: black !important;
        background-color: #ffcc00 !important;
        padding: 5px;
        border-radius: 5px;
        font-weight: bold;
    }
    .ok {
        color: white !important;
        background-color: #00cc66 !important;
        padding: 5px;
        border-radius: 5px;
        font-weight: bold;
    }
</style>
""", unsafe_allow_html=True)


# Sync from Google Sheets once after login to ensure fresh data
@st.cache_data(ttl=300, show_spinner="Syncing data from Google Sheets...")
def sync_data():
    db.pull_all_from_gsheets()

sync_data()

# Authentication Module
if 'logged_in' not in st.session_state:
    st.session_state['logged_in'] = False
if 'username' not in st.session_state:
    st.session_state['username'] = None
if 'role' not in st.session_state:
    st.session_state['role'] = None

if 'last_updated' not in st.session_state:
    st.session_state['last_updated'] = None

def mark_data_updated():
    st.session_state['last_updated'] = datetime.now()

if not st.session_state['logged_in']:
    st.title("🔒 Login")
    
    # Check if secret is detected to help user diagnose configuration issues
    if not db.INITIAL_ADMIN_PASSWORD:
        present_keys = []
        try:
            present_keys = list(st.secrets.keys())
        except:
            pass
        st.warning(f"⚠️ Configuration Warning: `INITIAL_ADMIN_PASSWORD` secret not detected. \n\n**Found keys:** {present_keys}. \n\nPlease ensure your secret name is exactly `INITIAL_ADMIN_PASSWORD` (all caps).")
    
    with st.form("login_form"):
        username = st.text_input("Username")
        password = st.text_input("Password", type="password")
        submitted = st.form_submit_button("Login")
        
        if submitted:
            # Clean inputs
            username = username.strip().lower()
            input_pwd = str(password).strip()
            
            # --- GUEST LOGIN BORDERCASE ---
            if username == 'guest':
                st.session_state['logged_in'] = True
                st.session_state['username'] = 'guest'
                st.session_state['role'] = 'Operator'
                st.success("Login successful (Guest Mode)!")
                st.rerun()

            try:
                # Ensure the admin exists right before login attempt just in case
                db.ensure_admin_exists()
                
                user_df = pd.read_sql_query("SELECT * FROM users WHERE LOWER(username) = ? AND password = ?", db.get_connection(), params=(username, input_pwd))
                
                # Check for successful DB login
                login_success = not user_df.empty
                
                # --- ADMIN RESCUE FALLBACK (Secrets) ---
                if not login_success and username == 'admin' and db.INITIAL_ADMIN_PASSWORD:
                    if input_pwd == str(db.INITIAL_ADMIN_PASSWORD).strip():
                        st.session_state['logged_in'] = True
                        st.session_state['username'] = 'admin'
                        st.session_state['role'] = 'Admin'
                        st.success("Login successful!")
                        st.rerun()
                
                if login_success:
                    st.session_state['logged_in'] = True
                    st.session_state['username'] = user_df.iloc[0]['username']
                    st.session_state['role'] = user_df.iloc[0]['role']
                    st.success("Login successful!")
                    st.rerun()
                else:
                    st.error("Invalid username or password.")
            except pd.errors.DatabaseError:
                st.error("Initializing database for the first time, please wait...")
                import init_db
                init_db.init_database()
                sync_data.clear()
                sync_data()
                st.rerun()

    st.markdown("---")
    if st.button("🔓 Continue as Guest (Standard Access)"):
        st.session_state['logged_in'] = True
        st.session_state['username'] = 'guest'
        st.session_state['role'] = 'Operator'
        st.success("Access Granted as Guest")
        st.rerun()
        
    st.stop() # Stop execution if not logged in



st.title("🏗️ EOT Crane Maintenance Tracking System")

col_info, col_inspection, col_drive, col_supervisor, col_sparing, col_logout = st.columns([2, 2, 2, 2, 2, 1])
with col_info:
    if st.session_state['last_updated']:
        st.markdown(f"**🕒 {st.session_state['last_updated'].strftime('%Y-%m-%d %H:%M:%S')}**")
with col_inspection:
    st.link_button("Safety Cell Check", "https://docs.google.com/document/d/1MF-F6MpADulb7lMhJUy1zxeHoEfV7ppc/edit?usp=sharing&ouid=103983074672171147767&rtpof=true&sd=true")
with col_drive:
    st.link_button("Officer Safety Check", "https://docs.google.com/spreadsheets/d/1EZqHb7qeCW8kkgCTNjaGdGLo5l-FeejY/edit?usp=sharing&ouid=103983074672171147767&rtpof=true&sd=true")
with col_supervisor:
    st.link_button("Supervisor Check", "https://docs.google.com/document/d/1IeCw0NXeJ83qLlyJv5aaAaAdshG2pqjZ/edit#heading=h.ena74ojzqd99")
with col_sparing:
    st.link_button("Request to Shop for Crane Sparing", "https://docs.google.com/document/d/1EeDaPC_Clp79c_kfLAYxkT4m09Pn2ydg2c46ZXnOt78/edit?usp=sharing")
# Note: Refresh Data button removed as per user request
with col_logout:
    if st.button("Logout"):
        st.session_state['logged_in'] = False
        st.session_state['username'] = None
        st.session_state['role'] = None
        st.rerun()

# Helper Functions
def load_data(table_name):
    query = f"SELECT * FROM {table_name}"
    return db.get_dataframe(query)

def generate_overdue_report(df):
    doc = Document()
    doc.add_heading('Overdue Maintenance Report', 0)
    
    # Add generation date and time
    now_str = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    doc.add_paragraph(f"Report Generated On: {now_str}")
    
    for c_type in ['Critical', 'Important', 'General']:
        doc.add_heading(f'{c_type} Cranes', level=1)
        # filter
        df_type = df[df['Crane Type'] == c_type]
        if df_type.empty:
            doc.add_paragraph('No overdue maintenance.')
            continue
            
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Crane ID'
        hdr_cells[1].text = 'Schedule Overdue'
        hdr_cells[2].text = 'Overdue Date'
        
        for _, row in df_type.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['crane_id'])
            row_cells[1].text = str(row['maintenance_type'])
            try:
                row_cells[2].text = pd.to_datetime(row['next_due_date']).strftime('%Y-%m-%d')
            except:
                row_cells[2].text = str(row['next_due_date'])
            
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def set_cell_background(cell, color):
    """Set the background color of a table cell (color is hex string like 'FFFF00')"""
    shading_elm = parse_xml(rf'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def generate_pivoted_maintenance_report(df_schedule, df_cranes):
    doc = Document()
    
    # Add Main Title and Generation Date
    doc.add_heading('Maintenance Schedule Report', 0)
    now_str = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    doc.add_paragraph(f"Report Generated On: {now_str}")
    
    # Merge to get crane type, location, and capacity
    df = df_schedule.merge(df_cranes[['id', 'type', 'location', 'capacity']], left_on='crane_id', right_on='id', how='left')
    df['Crane Type'] = df['type'].map({'A': 'Critical', 'B': 'Important', 'C': 'General'}).fillna('Unknown')
    
    # Define categories and their display names
    categories = [('Critical', 'Type A'), ('Important', 'Type B'), ('General', 'Type C')]
    
    for i, (cat_name, display_name) in enumerate(categories):
        df_cat = df[df['Crane Type'] == cat_name]
        if df_cat.empty:
            continue
            
        if i > 0:
            doc.add_page_break()
            
        doc.add_heading(f'Maintenance Schedule - {display_name} ({cat_name})', level=1)
        
        # Get unique schedule types and cranes for this category
        sched_types = sorted(df_cat['maintenance_type'].unique())
        crane_ids = sorted(df_cat['crane_id'].unique())
        
        # Create table: Crane No + Schedule Types
        table = doc.add_table(rows=1, cols=len(sched_types) + 1)
        table.style = 'Table Grid'
        
        # Header
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Crane No.'
        for j, s_type in enumerate(sched_types):
            hdr_cells[j+1].text = s_type
            
        for crane_id in crane_ids:
            row1 = table.add_row().cells
            row2 = table.add_row().cells
            
            # Get additional crane details
            info_row = df_cat[df_cat['crane_id'] == crane_id].iloc[0]
            location = info_row.get('location', 'N/A')
            capacity = info_row.get('capacity', 'N/A')
            
            row1[0].text = f"{crane_id}\n{location}\n({capacity})"
            row1[0].merge(row2[0])
            row1[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            for j, s_type in enumerate(sched_types):
                record = df_cat[(df_cat['crane_id'] == crane_id) & (df_cat['maintenance_type'] == s_type)]
                if not record.empty:
                    last_done = str(record.iloc[0]['last_maintenance_date'])
                    next_due_str = str(record.iloc[0]['next_due_date'])
                    status = record.iloc[0]['status']
                    
                    row1[j+1].text = f"Done: {last_done if last_done and last_done != 'None' else 'N/A'}"
                    
                    p = row2[j+1].paragraphs[0]
                    r1 = p.add_run("Due: ")
                    r2 = p.add_run(f"{next_due_str if next_due_str and next_due_str != 'None' else 'N/A'}")
                    
                    if status == 'Overdue':
                        r1.font.color.rgb = RGBColor(255, 0, 0)
                        r1.bold = True
                        r2.font.color.rgb = RGBColor(255, 0, 0)
                        r2.bold = True
                    elif status == 'Due Soon':
                        set_cell_background(row2[j+1], 'FFFF00')
                else:
                    row1[j+1].text = "-"
                    row2[j+1].text = "-"

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def save_data(df, table_name, index=False):
    db.save_dataframe(df, table_name, index)

# Evaluate Maintenance Status
def evaluate_status(date_val):
    if pd.isna(date_val):
        return "Unknown"
    
    try:
        date_obj = pd.to_datetime(date_val).date()
    except Exception:
        return "Unknown"
        
    today = datetime.now().date()
    delta = (date_obj - today).days
    
    if delta < 0:
        return "Overdue"
    elif delta <= 5:
        return "Due Soon"
    else:
        return "OK"

# Fetch Main Data
try:
    cranes_df = load_data("cranes")
    schedule_df = load_data("maintenance_schedule")
except Exception as e:
    st.error(f"Error loading data from database: {e}")
    st.info("Refreshing database from Google Sheets...")
    db.pull_all_from_gsheets()
    st.rerun()

# Recalculate status dynamically based on current date
schedule_df['status'] = schedule_df['next_due_date'].apply(evaluate_status)


# Tabs
is_admin = str(st.session_state.get('role', '')).lower() == 'admin'
is_guest = st.session_state.get('username', '') == 'guest'

if is_guest:
    tab_names = ["📊 Dashboard"]
else:
    tab_names = [
        "📊 Dashboard", 
        "🏗️ Cranes", 
        "📅 Maintenance Schedule", 
        "📖 Maintenance Log", 
        "🚨 Breakdown Log",
        "🧰 Spare Parts",
        "👤 Profile"
    ]
    if is_admin:
        tab_names.append("👥 Users")

tabs = st.tabs(tab_names)
tab1 = tabs[0]
if not is_guest:
    tab2 = tabs[1]
    tab3 = tabs[2]
    tab4 = tabs[3]
    tab5 = tabs[4]
    tab6 = tabs[5]
    tab7 = tabs[6]
    if is_admin:
        tab8 = tabs[7]

### ---------------- TAB 1: Dashboard ---------------- ###
with tab1:
    st.header("Dashboard Overview")
    
    # Filter for active cranes only for Dashboard
    active_cranes_df = cranes_df[cranes_df['status'] == 'Active']
    active_crane_ids = active_cranes_df['id'].tolist()
    active_schedule_df = schedule_df[schedule_df['crane_id'].isin(active_crane_ids)]
    
    # Logic: Filter overdue schedules
    raw_overdue_df = active_schedule_df[active_schedule_df['status'] == 'Overdue']
    
    # Prioritize the overdue schedule with the earliest due date (most overdue).
    # If there is a tie on the date, prioritize the higher schedule tier (S3 > S2 > S1).
    def get_tier(m_type):
        m_str = str(m_type)
        if len(m_str) >= 2 and m_str[0] == 'S' and m_str[1].isdigit():
            return int(m_str[1])
        return 0

    valid_overdue_indices = []
    
    if not raw_overdue_df.empty:
        temp_df = raw_overdue_df.copy()
        temp_df['due_date_dt'] = pd.to_datetime(temp_df['next_due_date'], errors='coerce')
        temp_df['tier'] = temp_df['maintenance_type'].apply(get_tier)
        
        # Sort so that for each crane, the first row is the highest schedule tier (S3 > S2 > S1)
        # and earliest date if tiers are equal
        temp_df = temp_df.sort_values(by=['crane_id', 'tier', 'due_date_dt'], ascending=[True, False, True])
        
        # Keep only the highest priority overdue schedule for each crane
        prioritized_df = temp_df.drop_duplicates(subset=['crane_id'], keep='first')
        
        valid_overdue_indices = prioritized_df.index.tolist()
        
    # Filter the overdue dataframe to only include the prioritized valid overdues
    overdue_df = raw_overdue_df.loc[valid_overdue_indices]

    # Compute Metrics
    total_cranes = len(active_cranes_df)
    
    due_this_week = len(active_schedule_df[active_schedule_df['status'] == 'Due Soon'])
    overdue = len(overdue_df)
    
    col1, col2, col3 = st.columns(3)
    col1.metric("Total Cranes", total_cranes)
    col2.metric("Maintenance Due (<=5 days)", due_this_week)
    col3.metric("Overdue Maintenance", overdue)
    
    st.divider()
    
    st.subheader("Metrics by Crane Type")
    # Base crane counts by type
    cranes_by_type = active_cranes_df['type'].value_counts() if 'type' in active_cranes_df.columns else pd.Series()
    type_a = cranes_by_type.get('A', 0)
    type_b = cranes_by_type.get('B', 0)
    type_c = cranes_by_type.get('C', 0)
    
    # Group valid overdue maintenance by type
    if 'type' in active_cranes_df.columns:
        overdue_merged = overdue_df.merge(active_cranes_df[['id', 'type']], left_on='crane_id', right_on='id', how='left')
        overdue_by_type = overdue_merged['type'].value_counts()
    else:
        overdue_by_type = pd.Series()
        
    od_a = overdue_by_type.get('A', 0)
    od_b = overdue_by_type.get('B', 0)
    od_c = overdue_by_type.get('C', 0)
    
    t_col1, t_col2, t_col3 = st.columns(3)
    t_col1.metric("Type A (Critical)", type_a, f"{od_a} Overdue", delta_color="inverse")
    t_col2.metric("Type B (Important)", type_b, f"{od_b} Overdue", delta_color="inverse")
    t_col3.metric("Type C (General)", type_c, f"{od_c} Overdue", delta_color="inverse")
    
    st.divider()
    
    st.subheader("Overdue Maintenance Details")
    # Prepare overdue data grouped by Crane Type and Schedule Type
    if not overdue_df.empty and 'type' in active_cranes_df.columns:
        od_merged = overdue_df.merge(active_cranes_df[['id', 'type']], left_on='crane_id', right_on='id', how='left')
        od_merged['Crane Type'] = od_merged['type'].map({'A': 'Critical', 'B': 'Important', 'C': 'General'}).fillna('Unknown')
        
        # Group by and aggregate counts and detailed strings for hover
        od_agg = od_merged.groupby(['Crane Type', 'maintenance_type']).apply(
            lambda x: pd.Series({
                'Count': len(x),
                'Details': '<br>'.join([f"Crane: {row['crane_id']} (Due: {pd.to_datetime(row['next_due_date']).strftime('%Y-%m-%d')})" for _, row in x.iterrows()])
            })
        ).reset_index()
        
        # Create three distinct charts
        ch_col1, ch_col2, ch_col3 = st.columns(3)
        
        def create_type_chart(c_type, title, color):
            df_type = od_agg[od_agg['Crane Type'] == c_type]
            if not df_type.empty:
                fig = px.bar(df_type, x='maintenance_type', y='Count', title=title, 
                             color_discrete_sequence=[color],
                             custom_data=['Details'])
                # Customize hover template
                fig.update_traces(hovertemplate="<b>Schedule:</b> %{x}<br><b>Count:</b> %{y}<br><br><b>Cranes:</b><br>%{customdata[0]}<extra></extra>")
                return fig
            else:
                return px.bar(title=f"{title} (None)", color_discrete_sequence=[color]) # Empty chart place holder
                
        with ch_col1:
            st.plotly_chart(create_type_chart('Critical', 'Type A (Critical)', '#ff4b4b'), use_container_width=True)
        with ch_col2:
            st.plotly_chart(create_type_chart('Important', 'Type B (Important)', '#ff9933'), use_container_width=True)
        with ch_col3:
            st.plotly_chart(create_type_chart('General', 'Type C (General)', '#3366cc'), use_container_width=True)

        # Added button here
        st.markdown("<br>", unsafe_allow_html=True)
        report_bytes = generate_overdue_report(od_merged)
        st.download_button(
            label="📄 Download Overdue Report (Word)",
            data=report_bytes,
            file_name=f"Overdue_Report_{datetime.now().strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
            
    else:
        st.info("No overdue maintenance found! ✅")
    
    st.divider()
    
    st.subheader("All Maintenance Statuses")
    if 'type' in active_cranes_df.columns:
        st_merged = active_schedule_df.merge(active_cranes_df[['id', 'type']], left_on='crane_id', right_on='id', how='left')
        st_merged['Crane Type'] = st_merged['type'].map({'A': 'Critical', 'B': 'Important', 'C': 'General'}).fillna('Unknown')
        sched_counts = st_merged.groupby(['Crane Type', 'maintenance_type', 'status']).size().reset_index(name='Count')
        
        def create_status_chart(c_type, title):
            df_type = sched_counts[sched_counts['Crane Type'] == c_type]
            if not df_type.empty:
                return px.bar(df_type, x='maintenance_type', y='Count', color='status', 
                              title=title,
                              color_discrete_map={"OK": "#00cc66", "Due Soon": "#ffcc00", "Overdue": "#ff4b4b"})
            else:
                return px.bar(title=f"{title} (None)") # Empty chart place holder
        
        st_col1, st_col2, st_col3 = st.columns(3)
        with st_col1:
            st.plotly_chart(create_status_chart('Critical', 'Type A (Critical) Statuses'), use_container_width=True)
        with st_col2:
            st.plotly_chart(create_status_chart('Important', 'Type B (Important) Statuses'), use_container_width=True)
        with st_col3:
            st.plotly_chart(create_status_chart('General', 'Type C (General) Statuses'), use_container_width=True)
    else:
        # Fallback if type is missing
        sched_counts = active_schedule_df.groupby(['maintenance_type', 'status']).size().reset_index(name='Count')
        fig2 = px.bar(sched_counts, x='maintenance_type', y='Count', color='status',
                      title="Status distribution across Schedules",
                      color_discrete_map={"OK": "#00cc66", "Due Soon": "#ffcc00", "Overdue": "#ff4b4b"})
        st.plotly_chart(fig2, use_container_width=True)

### ---------------- TAB 2: Cranes Master Database ---------------- ###
if not is_guest:
    with tab2:
        st.header("Crane Master Database")
        st.markdown("Edit crane details directly in the table below.")
        
        edited_cranes = st.data_editor(
            cranes_df,
            column_config={
                "status": st.column_config.SelectboxColumn(
                    "Status",
                    help="Operational status of the crane",
                    options=["Active", "Inactive"],
                    required=True,
                )
            },
            num_rows="dynamic" if is_admin else "fixed",
            use_container_width=True,
            key="cranes_editor",
            disabled=not is_admin
        )
        
        if not is_admin:
            st.info("⚠️ Only administrators can edit crane data.")
        
        if st.button("Save Crane Changes", type="primary", disabled=not is_admin):
            save_data(edited_cranes, "cranes")
            mark_data_updated()
            st.success("Crane Master Database updated successfully!")
            st.rerun()
        
### ---------------- TAB 3: Maintenance Schedule ---------------- ###
if not is_guest:
    with tab3:
        col_title_sched, col_report, col_sync = st.columns([2, 1, 1])
        with col_title_sched:
            st.header("Maintenance Schedule")
        with col_report:
            st.markdown("<br>", unsafe_allow_html=True)
            report_pivoted = generate_pivoted_maintenance_report(schedule_df, cranes_df)
            st.download_button(
                label="📄 Download Schedule Report",
                data=report_pivoted,
                file_name=f"Maintenance_Schedule_Report_{datetime.now().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_sched_report"
            )
        with col_sync:
            if st.button("🔄 Sync with Schedule Master", disabled=not is_admin):
                import db as db_local
                try:
                    sm = load_data("Schedule_Master")
                    # Ensure Frequency column can be parsed to int
                    sm['Frequency'] = pd.to_numeric(sm['Frequency'], errors='coerce').fillna(30)
                    freq_dict = dict(zip(sm['Schedule'], sm['Frequency']))
                except Exception as e:
                    st.error(f"Could not load Schedule_Master from Google Sheets: {e}")
                    freq_dict = {}
                
                # Recalculate with cascading hierarchy logic
                df_to_update = load_data("maintenance_schedule")
                df_to_update['last_m_dt'] = pd.to_datetime(df_to_update['last_maintenance_date'], errors='coerce')
                
                for crane in df_to_update['crane_id'].unique():
                    crane_scheds = df_to_update[df_to_update['crane_id'] == crane]
                    for base in ['A', 'B', 'C']:
                        s1_row = crane_scheds[crane_scheds['maintenance_type'] == f'S1{base}']
                        s2_row = crane_scheds[crane_scheds['maintenance_type'] == f'S2{base}']
                        s3_row = crane_scheds[crane_scheds['maintenance_type'] == f'S3{base}']
                        
                        s1_dt = s1_row['last_m_dt'].max() if not s1_row.empty else pd.NaT
                        s2_dt = s2_row['last_m_dt'].max() if not s2_row.empty else pd.NaT
                        s3_dt = s3_row['last_m_dt'].max() if not s3_row.empty else pd.NaT
                        
                        eff_s1 = max([d for d in [s1_dt, s2_dt, s3_dt] if pd.notna(d)], default=pd.NaT)
                        eff_s2 = max([d for d in [s2_dt, s3_dt] if pd.notna(d)], default=pd.NaT)
                        eff_s3 = max([d for d in [s3_dt] if pd.notna(d)], default=pd.NaT)
                        
                        if not s1_row.empty and pd.notna(eff_s1):
                            df_to_update.loc[s1_row.index, 'eff_last'] = eff_s1
                        if not s2_row.empty and pd.notna(eff_s2):
                            df_to_update.loc[s2_row.index, 'eff_last'] = eff_s2
                        if not s3_row.empty and pd.notna(eff_s3):
                            df_to_update.loc[s3_row.index, 'eff_last'] = eff_s3
    
                queries = []
                for _, row in df_to_update.iterrows():
                    try:
                        last_m = row.get('eff_last')
                        if pd.notna(last_m):
                            m_type = row['maintenance_type']
                            days_add = freq_dict.get(m_type, 30)
                            next_due = (last_m + pd.Timedelta(days=days_add)).strftime('%Y-%m-%d')
                            queries.append((last_m.strftime('%Y-%m-%d'), next_due, row['id']))
                    except:
                        pass
                
                if queries:
                    db_local.execute_many_query("UPDATE maintenance_schedule SET last_maintenance_date = ?, next_due_date = ?, status='OK' WHERE id = ?", queries)
                    mark_data_updated()
                    st.success("Synchronized successfully!")
                st.rerun()
        
        # Merge crane type into schedule df
        schedule_with_type = schedule_df.merge(cranes_df[['id', 'type']], left_on='crane_id', right_on='id', how='left')
        schedule_with_type['Crane Type'] = schedule_with_type['type'].map({'A': 'Critical', 'B': 'Important', 'C': 'General'}).fillna('Unknown')
    
        # Filtering
        col_f1, col_f2, col_f3, col_f4 = st.columns(4)
        with col_f1:
            f_crane_id = st.text_input("Search Crane No:", placeholder="e.g., MW-140")
        with col_f2:
            f_status = st.multiselect("Filter by Status:", options=["OK", "Due Soon", "Overdue"], default=["OK", "Due Soon", "Overdue"])
        with col_f3:
            f_type = st.multiselect("Filter Interval:", options=schedule_df['maintenance_type'].unique(), default=schedule_df['maintenance_type'].unique())
        with col_f4:
            f_crane_type = st.multiselect("Filter Category:", options=["Critical", "Important", "General", "Unknown"], default=["Critical", "Important", "General", "Unknown"])
        
        filtered_schedule = schedule_with_type[
            (schedule_with_type['status'].isin(f_status)) & 
            (schedule_with_type['maintenance_type'].isin(f_type)) &
            (schedule_with_type['Crane Type'].isin(f_crane_type))
        ]
        
        if f_crane_id:
            # Case insensitive exact or partial match
            filtered_schedule = filtered_schedule[filtered_schedule['crane_id'].str.contains(f_crane_id, case=False, na=False)]
            
        filtered_schedule = filtered_schedule.drop(columns=['id', 'id_x', 'id_y', 'type'], errors='ignore') # Drop merge artifacts
        
        # Custom styling function for dataframe display
        def color_status(val):
            color = ''
            if val == 'Overdue':
                color = 'background-color: #ff4b4b; color: white'
            elif val == 'Due Soon':
                color = 'background-color: #ffcc00; color: black'
            elif val == 'OK':
                color = 'background-color: #00cc66; color: white'
            return color
    
        st.dataframe(
            filtered_schedule.style.applymap(color_status, subset=['status']),
            use_container_width=True,
            hide_index=True
        )
    
@st.cache_data
def get_schedule_frequencies():
    try:
        sm = db.get_dataframe("SELECT * FROM Schedule_Master")
        # Ensure Frequency column can be parsed to int
        sm['Frequency'] = pd.to_numeric(sm['Frequency'], errors='coerce').fillna(30)
        return dict(zip(sm['Schedule'], sm['Frequency']))
    except Exception as e:
        return {}

schedule_freq_map = get_schedule_frequencies()

# Centralized Save Logic for Maintenance Log
def commit_maintenance_log(log_data):
    dt_taking = log_data['taking_over'].strftime('%Y-%m-%d %H:%M:%S')
    dt_handing = log_data['handing_over'].strftime('%Y-%m-%d %H:%M:%S')
    l_date = log_data['date']
    l_crane = log_data['crane']
    l_type = log_data['type']
    l_status = log_data['status']
    l_remarks = log_data['remarks']
    l_photo = log_data['photo']
    
    try:
        max_id_df = db.get_dataframe("SELECT MAX(CAST(id AS INTEGER)) as max_id FROM maintenance_logs")
        new_id = int(max_id_df.iloc[0]['max_id']) + 1 if pd.notna(max_id_df.iloc[0]['max_id']) else 1
    except:
        new_id = 1
    
    if l_photo:
        with st.spinner("Uploading photograph..."):
            photo_path = db.upload_image_to_drive(l_photo)
    else:
        photo_path = ""
    
    db.execute_query("""
        INSERT INTO maintenance_logs (id, date, crane_id, maintenance_type, taking_over_datetime, handing_over_datetime, checklist_status, remarks, photo_path)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (new_id, l_date.strftime('%Y-%m-%d'), l_crane, l_type, dt_taking, dt_handing, l_status, l_remarks, photo_path))
    
    # Update cascading schedules
    types_to_update = [l_type]
    if l_type.startswith('S3'):
        cl = l_type[2:] if len(l_type) > 2 else ''
        types_to_update.extend([f"S2{cl}", f"S1{cl}"])
    elif l_type.startswith('S2'):
        cl = l_type[2:] if len(l_type) > 2 else ''
        types_to_update.extend([f"S1{cl}"])
    
    for t in types_to_update:
        t_days_add = schedule_freq_map.get(t, 30)
        t_next_due = l_date + pd.Timedelta(days=t_days_add)
        db.execute_query("""
            UPDATE maintenance_schedule SET last_maintenance_date = ?, next_due_date = ?, status = 'OK'
            WHERE crane_id = ? AND maintenance_type = ?
        """, (l_date.strftime('%Y-%m-%d'), t_next_due.strftime('%Y-%m-%d'), l_crane, t))
    
    st.success("Entry saved successfully!")
    mark_data_updated()
    st.rerun()

# Modal Dialog for Confirmation
@st.dialog("⚠️ Verify Maintenance Entry")
def confirm_maint_dialog(log_data):
    st.markdown("Please review all details before final submission.")
    
    col1, col2 = st.columns(2)
    with col1:
        st.write(f"**Crane ID:** {log_data['crane']}")
        st.write(f"**Maintenance Date:** {log_data['date'].strftime('%Y-%m-%d')}")
        st.write(f"**Taking Over:** {log_data['taking_over'].strftime('%Y-%m-%d %H:%M')}")
    with col2:
        st.write(f"**Schedule:** {log_data['type']}")
        st.write(f"**Overall Status:** {log_data['status']}")
        st.write(f"**Handing Over:** {log_data['handing_over'].strftime('%Y-%m-%d %H:%M')}")
    
    if log_data['remarks']:
        st.info(f"**Remarks:** {log_data['remarks']}")
    else:
        st.write("**Remarks:** None provided.")
    
    st.divider()
    c1, c2 = st.columns(2)
    if c1.button("✅ CONFIRM & SAVE", type="primary", use_container_width=True):
        commit_maintenance_log(log_data)
    if c2.button("❌ CANCEL", use_container_width=True):
        st.rerun()


### ---------------- TAB 4: Maintenance Log ---------------- ###
if not is_guest:
    with tab4:
        st.header("Maintenance Log")
        
        colA, colB = st.columns([1, 2])
        
        with colA:
            st.subheader("Add New Entry")
            l_crane = st.selectbox("Crane ID *", cranes_df['id'].unique(), key='l_crane_maint')
            l_available_schedules = schedule_df[schedule_df['crane_id'] == l_crane]['maintenance_type'].unique()
            
            if len(l_available_schedules) == 0:
                c_type_row = cranes_df[cranes_df['id'] == l_crane]
                c_type_val = c_type_row.iloc[0]['type'] if not c_type_row.empty else 'A'
                l_available_schedules = [f"S1{c_type_val}", f"S2{c_type_val}", f"S3{c_type_val}"]
                
            with st.form("log_form", clear_on_submit=False):
                l_date = st.date_input("Maintenance Date *")
                l_type = st.selectbox("Maintenance Schedule *", l_available_schedules)
                
                col_t1, col_t2 = st.columns(2)
                with col_t1:
                    l_taking_over_d = st.date_input("Taking Over Date *")
                    l_taking_over_t = st.time_input("Taking Over Time *")
                with col_t2:
                    l_handing_over_d = st.date_input("Handing Over Date *")
                    l_handing_over_t = st.time_input("Handing Over Time *")
                    
                l_status = st.selectbox("Overall Checklist Status *", ['Completed OK', 'Pending Action', 'Failed'])
                l_remarks = st.text_area("Remarks (Optional)")
                l_photo = st.file_uploader("Upload Document/Photo (Optional)", type=['jpg', 'png', 'jpeg', 'pdf'])
                
                if not is_admin:
                    st.info("⚠️ Only administrators can submit logs.")
                
                submitted = st.form_submit_button("Submit Log", disabled=not is_admin)
                
                if submitted:
                    log_data = {
                        'date': l_date,
                        'crane': l_crane,
                        'type': l_type,
                        'taking_over': datetime.combine(l_taking_over_d, l_taking_over_t),
                        'handing_over': datetime.combine(l_handing_over_d, l_handing_over_t),
                        'status': l_status,
                        'remarks': l_remarks,
                        'photo': l_photo
                    }
                    confirm_maint_dialog(log_data)
    
        with colB:
            st.subheader("Historical Logs")
            logs_df = load_data("maintenance_logs")
            
            # Process photo_path to extract link and display text
            if not logs_df.empty and 'photo_path' in logs_df.columns:
                def extract_display_text(val):
                    if "|" in str(val):
                        return val.split("|")[0]
                    return "View File" if str(val).startswith("http") else val
                
                def extract_url(val):
                    if "|" in str(val):
                        return val.split("|")[1]
                    return val if str(val).startswith("http") else ""
    
                logs_df['Link'] = logs_df['photo_path'].apply(extract_url)
                logs_df['Document Name'] = logs_df['photo_path'].apply(extract_display_text)
                
                # Reorder or drop original
                cols = [c for c in logs_df.columns if c != 'photo_path']
                logs_df = logs_df[cols]
    
            st.dataframe(
                logs_df, 
                use_container_width=True, 
                hide_index=True,
                column_config={
                    "Link": st.column_config.LinkColumn(
                        "Download", 
                        help="Click to download",
                        display_text="Download PDF"
                    )
                }
            )

### ---------------- TAB 5: Breakdown Log ---------------- ###
if not is_guest:
    with tab5:
        st.header("Breakdown Log")
        
        colA, colB = st.columns([1, 2])
        
        with colA:
            st.subheader("Add New Entry")
            
            # Display selected Crane ID outside the form for dynamic reloading
            b_crane = st.selectbox("Crane ID", cranes_df['id'].unique(), key='b_crane_breakdown')
            
            c_scheds = schedule_df[schedule_df['crane_id'] == b_crane]
            
            st.markdown("**Current Schedule Info:**")
            cols_s = st.columns(3)
            
            def get_sched_text(prefix):
                df_s = c_scheds[c_scheds['maintenance_type'].str.startswith(prefix)]
                if df_s.empty:
                    return f"{prefix} Not Found\nLast: N/A\nNext: N/A"
                # Prioritize higher letter? Actually typically there is only 1 active type per prefix per crane, e.g. S1A
                row = df_s.iloc[0]
                return f"Type: {row['maintenance_type']}\nLast: {row['last_maintenance_date']}\nNext: {row['next_due_date']}"
    
            with cols_s[0]:
                st.text_area("S1 Schedule", get_sched_text('S1'), disabled=True, height=105)
            with cols_s[1]:
                st.text_area("S2 Schedule", get_sched_text('S2'), disabled=True, height=105)
            with cols_s[2]:
                st.text_area("S3 Schedule", get_sched_text('S3'), disabled=True, height=105)
                
            try:
                assemblies_df = load_data("failure_assemblies")
                assembly_options = assemblies_df['assembly_name'].tolist()
            except:
                assembly_options = ["LT Assembly", "CT Assembly", "Other"]
                
            b_failure_assembly = st.selectbox("Failure of Assembly", assembly_options)
            
            try:
                comp_df = pd.read_sql_query("SELECT component_name FROM failure_components WHERE assembly_name = ?", db.get_connection(), params=(b_failure_assembly,))
                comp_options = comp_df['component_name'].tolist() if not comp_df.empty else ["Other"]
            except:
                comp_options = ["Other"]
            b_failure_component = st.selectbox("Child Component", comp_options)
    
            try:
                def_df = pd.read_sql_query("SELECT defect_name FROM failure_defects WHERE component_name = ?", db.get_connection(), params=(b_failure_component,))
                def_options = def_df['defect_name'].tolist() if not def_df.empty else ["Other"]
            except:
                def_options = ["Other"]
            b_failure_defect = st.selectbox("Nature of Defect", def_options)
    
            with st.form("breakdown_form"):
                col_rep1, col_rep2 = st.columns(2)
                with col_rep1:
                    b_reported_d = st.date_input("Breakdown Reported Date", key="b_rep_d")
                    b_reported_t = st.time_input("Breakdown Reported Time", key="b_rep_t")
                with col_rep2:
                    # Spacer or empty to match design
                    pass
                    
                col_b1, col_b2 = st.columns(2)
                with col_b1:
                    b_taking_over_d = st.date_input("Taking Over Date", key="b_tak_d")
                    b_taking_over_t = st.time_input("Taking Over Time", key="b_tak_t")
                with col_b2:
                    b_handing_over_d = st.date_input("Handing Over Date", key="b_hand_d")
                    b_handing_over_t = st.time_input("Handing Over Time", key="b_hand_t")
                    
                b_reported_failure = st.text_input("Reported Failure Type")
                b_root_cause = st.text_input("Root Cause of Failure")
                b_corrective_action = st.text_area("Corrective Action Taken")
                    
                b_status = st.selectbox("Overall Checklist Status", ['Completed OK', 'Pending Action', 'Failed'], key="b_status")
                b_remarks = st.text_area("Remarks", key="b_remarks")
                b_photo = st.file_uploader("Upload Document/Photo", type=['jpg', 'png', 'jpeg', 'pdf'], key="b_photo")
                if not is_admin:
                    st.info("⚠️ Only administrators can submit logs.")
                b_submitted = st.form_submit_button("Submit Breakdown Log", disabled=not is_admin)
                
                if b_submitted:
                    dt_reported = datetime.combine(b_reported_d, b_reported_t).strftime('%Y-%m-%d %H:%M:%S')
                    dt_taking = datetime.combine(b_taking_over_d, b_taking_over_t).strftime('%Y-%m-%d %H:%M:%S')
                    dt_handing = datetime.combine(b_handing_over_d, b_handing_over_t).strftime('%Y-%m-%d %H:%M:%S')
                    # Get dynamic ID for breakdown_logs
                    try:
                        max_id_df = db.get_dataframe("SELECT MAX(CAST(id AS INTEGER)) as max_id FROM breakdown_logs")
                        new_id = int(max_id_df.iloc[0]['max_id']) + 1 if pd.notna(max_id_df.iloc[0]['max_id']) else 1
                    except:
                        new_id = 1
                        
                    # Insert into DB
                    if b_photo:
                        with st.spinner("Uploading photograph to Google Drive..."):
                            b_photo_path = db.upload_image_to_drive(b_photo)
                    else:
                        b_photo_path = ""
                    db.execute_query("""
                        INSERT INTO breakdown_logs (
                            id, crane_id, breakdown_reported_datetime, taking_over_datetime, handing_over_datetime, 
                            checklist_status, remarks, photo_path, failure_assembly, reported_failure_type, 
                            root_cause_failure, corrective_action, failure_component, failure_defect
                        )
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """, (new_id, b_crane, dt_reported, dt_taking, dt_handing, 
                          b_status, b_remarks, b_photo_path, b_failure_assembly, b_reported_failure, 
                          b_root_cause, b_corrective_action, b_failure_component, b_failure_defect))
                    
                    st.success("Breakdown log added!")
                    mark_data_updated()
                    st.rerun()
    
        with colB:
            st.subheader("Historical Breakdown Logs")
            try:
                b_logs_df = load_data("breakdown_logs")
                
                if not b_logs_df.empty and 'photo_path' in b_logs_df.columns:
                    def extract_display_text(val):
                        if "|" in str(val):
                            return val.split("|")[0]
                        return "View File" if str(val).startswith("http") else val
                    
                    def extract_url(val):
                        if "|" in str(val):
                            return val.split("|")[1]
                        return val if str(val).startswith("http") else ""
    
                    b_logs_df['Link'] = b_logs_df['photo_path'].apply(extract_url)
                    b_logs_df['Document Name'] = b_logs_df['photo_path'].apply(extract_display_text)
                    
                    cols = [c for c in b_logs_df.columns if c != 'photo_path']
                    b_logs_df = b_logs_df[cols]
    
                st.dataframe(
                    b_logs_df, 
                    use_container_width=True, 
                    hide_index=True,
                    column_config={
                        "Link": st.column_config.LinkColumn(
                            "Download", 
                            help="Click to download",
                            display_text="Download PDF"
                        )
                    }
                )
            except Exception:
                st.info("No breakdown logs found.")

### ---------------- TAB 6: Spare Parts Inventory ---------------- ###
if not is_guest:
    with tab6:
        st.header("Spare Parts Inventory")
        
        colA, colB = st.columns([1, 2])
        
        with colA:
            st.subheader("Add New Spare Part")
            with st.form("spare_part_form"):
                sp_name = st.text_input("Part Name")
                
                all_cranes = cranes_df['id'].unique().tolist()
                sp_cranes = st.multiselect("Applicable Cranes", all_cranes, help="Select one or multiple cranes that this part applies to.")
                
                col_sp1, col_sp2 = st.columns(2)
                with col_sp1:
                    sp_stock = st.number_input("Stock Quantity", min_value=0, step=1)
                with col_sp2:
                    sp_min_stock = st.number_input("Minimum Stock", min_value=0, step=1)
                    
                sp_supplier = st.text_input("Supplier")
                sp_last_replacement = st.date_input("Last Replacement Date")
                sp_remarks = st.text_area("Remarks")
                if not is_admin:
                    st.info("⚠️ Only administrators can add inventory.")
                sp_submitted = st.form_submit_button("Add Spare Part", disabled=not is_admin)
                if sp_submitted:
                    # Get dynamic ID for spare_parts
                    try:
                        max_id_df = db.get_dataframe("SELECT MAX(CAST(id AS INTEGER)) as max_id FROM spare_parts")
                        new_id = int(max_id_df.iloc[0]['max_id']) + 1 if pd.notna(max_id_df.iloc[0]['max_id']) else 1
                    except:
                        new_id = 1
                        
                    cranes_str = ", ".join(sp_cranes) if sp_cranes else "None"
                    db.execute_query("""
                        INSERT INTO spare_parts (
                            id, part_name, applicable_cranes, stock_quantity, minimum_stock, supplier, last_replacement_date, remarks
                        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                    """, (new_id, sp_name, cranes_str, sp_stock, sp_min_stock, sp_supplier, sp_last_replacement.strftime('%Y-%m-%d'), sp_remarks))
                    st.success(f"Added {sp_name} to inventory!")
                    mark_data_updated()
                    st.rerun()
    
        with colB:
            st.subheader("Current Inventory")
            parts_df = load_data("spare_parts")
            
            if parts_df.empty:
                st.info("No spare parts currently logged in inventory.")
                parts_df = pd.DataFrame(columns=['id', 'part_name', 'applicable_cranes', 'stock_quantity', 'minimum_stock', 'supplier', 'last_replacement_date'])
    
            edited_parts = st.data_editor(
                parts_df,
                num_rows="dynamic" if is_admin else "fixed",
                use_container_width=True,
                key="parts_editor",
                disabled=not is_admin
            )
            
            if not is_admin:
                st.info("⚠️ Only administrators can edit inventory levels.")
                
            if st.button("Save Edits to Inventory", type="primary", disabled=not is_admin):
                save_data(edited_parts, "spare_parts")
                mark_data_updated()
                st.success("Spare Parts Inventory updated successfully!")
                st.rerun()

### ---------------- TAB 7: Profile (Change Password) ---------------- ###
if not is_guest:
    with tab7:
        st.header("👤 Profile")
        st.subheader("Change My Password")
        with st.form("change_password_form"):
            old_pass = st.text_input("Current Password", type="password")
            new_pass = st.text_input("New Password", type="password")
            confirm_pass = st.text_input("Confirm New Password", type="password")
            submitted_pwd = st.form_submit_button("Change Password")
            
            if submitted_pwd:
                if not old_pass or not new_pass or not confirm_pass:
                    st.warning("All fields are required.")
                elif new_pass != confirm_pass:
                    st.error("New passwords do not match.")
                else:
                    user_df = pd.read_sql_query("SELECT * FROM users WHERE username = ? AND password = ?", db.get_connection(), params=(st.session_state['username'], old_pass))
                    if user_df.empty:
                        st.error("Incorrect current password.")
                    else:
                        db.execute_query("UPDATE users SET password = ? WHERE username = ?", (new_pass, st.session_state['username']))
                        mark_data_updated()
                        st.success("Password changed successfully!")
                        st.rerun()

### ---------------- TAB 8: Users (Admin Only) ---------------- ###
if is_admin:
    with tab8:
        st.header("👥 User Management")
        
        users_df = load_data("users")
        
        col_u1, col_u2 = st.columns([1, 2])
        with col_u1:
            st.subheader("Add New User")
            with st.form("new_user_form"):
                new_username = st.text_input("Username")
                new_password = st.text_input("Password", type="password")
                new_role = st.selectbox("Role", ["Operator", "Admin"])
                
                submitted_user = st.form_submit_button("Create User")
                if submitted_user:
                    if not new_username or not new_password:
                        st.error("Username and password are required")
                    else:
                        try:
                            db.execute_query("INSERT INTO users (username, password, role) VALUES (?, ?, ?)", (new_username, new_password, new_role))
                            st.success(f"User '{new_username}' created successfully!")
                            mark_data_updated()
                            st.rerun()
                        except Exception as e:
                            st.error("Failed to create user. Username might already exist.")
                            
        with col_u2:
            st.subheader("Manage Users")
            st.dataframe(users_df[['id', 'username', 'role']], use_container_width=True, hide_index=True)
            
            st.markdown("---")
            st.subheader("Delete User")
            
            # Don't let the currently logged in admin delete themselves
            deletable_users = users_df[users_df['username'] != st.session_state['username']]['username'].tolist()
            with st.form("delete_user_form"):
                del_user = st.selectbox("Select User to Delete", deletable_users if deletable_users else ["None Available"])
                submitted_del = st.form_submit_button("Delete Selected User", type="primary")
                if submitted_del:
                    if del_user == "None Available":
                        st.warning("No user selected.")
                    else:
                        db.execute_query("DELETE FROM users WHERE username = ?", (del_user,))
                        mark_data_updated()
                        st.success(f"User '{del_user}' deleted!")
                        st.rerun()
